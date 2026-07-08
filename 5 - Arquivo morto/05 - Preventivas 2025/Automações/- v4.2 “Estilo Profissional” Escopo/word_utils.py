# word_utils.py — (mesmo da sua versão estável; incluído aqui completo para manter “envio sempre completo”)
import os
import datetime
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK
from PIL import Image, UnidentifiedImageError

PASTAS_TEXTO_NORMAL = ["- Detalhes", "- Vista ampla"]


def aplicar_estilo(run, tamanho, negrito=False):
    run.font.name = "Arial"
    run.font.size = Pt(tamanho)
    run.bold = negrito


def inserir_conteudo(modelo_path, conteudo, output_path):
    doc = Document(modelo_path)
    contador_imagens = 0
    paragrafo_insercao_index = None

    for i, paragrafo in enumerate(doc.paragraphs):
        if "{{start_here}}" in paragrafo.text:
            paragrafo_insercao_index = i
            break

    if paragrafo_insercao_index is None:
        print("Marca '{{start_here}}' não encontrada no modelo.")
        return contador_imagens

    # Caminhos de log
    base_dir = os.path.dirname(output_path)
    log_falhas = os.path.join(base_dir, "imagens_falhas.txt")
    log_nao_inseridas = os.path.join(base_dir, "imagens_nao_inseridas.txt")

    with open(log_falhas, "w", encoding="utf-8") as f:
        f.write(f"Log de imagens com falha - {datetime.datetime.now()}\n\n")
    with open(log_nao_inseridas, "w", encoding="utf-8") as f:
        f.write(f"Imagens não inseridas - {datetime.datetime.now()}\n\n")

    imagens_recebidas = [item["imagem"] for item in conteudo if isinstance(item, dict) and "imagem" in item]
    imagens_inseridas = []

    conteudo_invertido = list(reversed(conteudo))

    for item in conteudo_invertido:
        if isinstance(item, str):
            titulo = item.replace("»", "").strip() + ":"
            nivel = item.count("»")

            p = doc.paragraphs[paragrafo_insercao_index].insert_paragraph_before('')
            run = p.add_run(titulo)

            if any(pasta in titulo for pasta in PASTAS_TEXTO_NORMAL):
                aplicar_estilo(run, 11, negrito=True)
                p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            elif nivel == 0:
                p.style = 'Heading 1'
            elif nivel == 1:
                p.style = 'Heading 2'
            elif nivel == 2:
                p.style = 'Heading 3'
            else:
                aplicar_estilo(run, 12, negrito=True)

        elif isinstance(item, dict) and "imagem" in item:
            imagem_path = item["imagem"]

            if os.path.exists(imagem_path) and os.path.getsize(imagem_path) > 0:
                try:
                    p = doc.paragraphs[paragrafo_insercao_index].insert_paragraph_before('')
                    with Image.open(imagem_path) as img:
                        largura_original, altura_original = img.size

                        # proporção segura
                        altura_desejada_cm = 10
                        if altura_original > 0:
                            proporcao = altura_desejada_cm / altura_original * 2.54
                            largura_proporcional_cm = largura_original * proporcao / 2.54
                        else:
                            proporcao = None
                            largura_proporcional_cm = None

                        run = p.add_run()
                        try:
                            # tentativa com redimensionamento
                            if largura_proporcional_cm and largura_proporcional_cm > 0:
                                run.add_picture(
                                    imagem_path,
                                    width=Cm(largura_proporcional_cm),
                                    height=Cm(altura_desejada_cm)
                                )
                            else:
                                raise ValueError("Dimensão inválida calculada")

                        except Exception:
                            # fallback sem redimensionamento
                            run.add_picture(imagem_path)

                        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        contador_imagens += 1
                        imagens_inseridas.append(imagem_path)

                except UnidentifiedImageError:
                    msg = f"[ERRO: Formato de imagem não reconhecido] {imagem_path}"
                    print(msg)
                    p = doc.paragraphs[paragrafo_insercao_index].insert_paragraph_before(msg)
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    with open(log_falhas, "a", encoding="utf-8") as f:
                        f.write(msg + "\n")

                except Exception as e:
                    msg = f"[ERRO ao inserir imagem: {imagem_path}] Detalhe: {e}"
                    print(msg)
                    p = doc.paragraphs[paragrafo_insercao_index].insert_paragraph_before(msg)
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    with open(log_falhas, "a", encoding="utf-8") as f:
                        f.write(msg + "\n")
            else:
                msg = f"[ERRO: Arquivo de imagem inválido] {imagem_path}"
                print(msg)
                p = doc.paragraphs[paragrafo_insercao_index].insert_paragraph_before(msg)
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                with open(log_falhas, "a", encoding="utf-8") as f:
                    f.write(msg + "\n")

        elif isinstance(item, dict) and "quebra_pagina" in item:
            doc.paragraphs[paragrafo_insercao_index].insert_paragraph_before('').add_run().add_break(WD_BREAK.PAGE)

    # Salva documento
    doc.save(output_path)

    # Gera log de imagens não inseridas
    nao_inseridas = [img for img in imagens_recebidas if img not in imagens_inseridas]
    with open(log_nao_inseridas, "a", encoding="utf-8") as f:
        for img in nao_inseridas:
            f.write(img + "\n")

    return contador_imagens
