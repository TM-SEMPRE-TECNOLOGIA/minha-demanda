import os
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image, UnidentifiedImageError

# =========================
# CONFIG
# =========================
PASTAS_TEXTO_NORMAL = ["- Detalhes", "- Vista ampla"]

ALTURA_PADRAO = 10.0  # cm
LARGURA_MAX_3_COL = 5.63  # cm
LARGURA_MAX_2_COL = 7.50  # cm

LABEL_TEXT = "0,00 m"

# Label sizing (pt) — sized to avoid truncation and allow edits like "21,05 m"
LABEL_W_PT = 90.0
LABEL_H_PT = 18.0
LABEL_PAD_PT = 6.0

# Border: 1.5pt in DrawingML uses 1pt = 12700 EMU
LINE_W_EMU = int(1.5 * 12700)  # 19050


# =========================
# UNIT CONVERSIONS
# =========================
def _cm_to_pt(cm: float) -> float:
    return cm * 28.3464567

def _pt_to_emu(pt: float) -> int:
    # 1 inch = 72pt = 914400 EMU => 1pt = 12700 EMU
    return int(pt * 12700)

def _cm_to_emu(cm: float) -> int:
    # 1 cm = 360000 EMU
    return int(cm * 360000)


# =========================
# SKIP RULES
# =========================
def _should_skip_labels(path: str) -> bool:
    p = (path or "").lower()
    return ("- detalhes" in p) or ("- vista ampla" in p)


# =========================
# IMAGE ANALYSIS / LAYOUT GROUPING
# =========================
def analisar_imagem(caminho_imagem):
    try:
        with Image.open(caminho_imagem) as img:
            w, h = img.size
            aspect_ratio = w / h
            largura_em_10cm = ALTURA_PADRAO * aspect_ratio
            is_vertical = h > w
            return is_vertical, largura_em_10cm
    except Exception:
        return False, 999.0


def otimizar_layout(conteudo_original):
    novo_conteudo = []
    buffer_imagens = []

    def processar_buffer():
        nonlocal buffer_imagens
        while buffer_imagens:
            qtd = len(buffer_imagens)
            grupo_formado = False

            if qtd >= 3:
                candidatos = buffer_imagens[:3]
                if all(img['largura_10cm'] <= LARGURA_MAX_3_COL for img in candidatos):
                    novo_conteudo.append({
                        "tabela_imagens": [img['caminho'] for img in candidatos],
                        "colunas": 3
                    })
                    buffer_imagens = buffer_imagens[3:]
                    grupo_formado = True

            if not grupo_formado and qtd >= 2:
                candidatos = buffer_imagens[:2]
                if all(img['largura_10cm'] <= LARGURA_MAX_2_COL for img in candidatos):
                    novo_conteudo.append({
                        "tabela_imagens": [img['caminho'] for img in candidatos],
                        "colunas": 2
                    })
                    buffer_imagens = buffer_imagens[2:]
                    grupo_formado = True

            if not grupo_formado:
                img = buffer_imagens.pop(0)
                novo_conteudo.append({"imagem": img['caminho']})

    for item in conteudo_original:
        if isinstance(item, dict) and "imagem" in item:
            caminho = item["imagem"]
            is_vert, larg_10cm = analisar_imagem(caminho)

            if is_vert:
                buffer_imagens.append({
                    "caminho": caminho,
                    "largura_10cm": larg_10cm
                })
            else:
                processar_buffer()
                novo_conteudo.append(item)
        else:
            processar_buffer()
            novo_conteudo.append(item)

    processar_buffer()
    return novo_conteudo


def aplicar_estilo(run, size_pt, is_bold=False):
    run.font.size = Pt(size_pt)
    run.font.bold = is_bold


# =========================
# INLINE -> ANCHOR (FLOATING)
# =========================
def _convert_last_inline_to_anchor(run):
    """
    Converte o último wp:inline do run em wp:anchor (imagem flutuante).
    Isso permite sobreposição estável de shapes (DrawingML).
    """
    drawing_list = run._r.xpath('./w:drawing')
    if not drawing_list:
        return
    drawing = drawing_list[0]
    inline_list = drawing.xpath('./wp:inline')
    if not inline_list:
        return
    wp_inline = inline_list[0]

    anchor = OxmlElement('wp:anchor')
    anchor.set(qn('wp:simplePos'), '0')
    anchor.set(qn('wp:relativeHeight'), '251659264')
    anchor.set(qn('wp:behindDoc'), '0')
    anchor.set(qn('wp:locked'), '0')
    anchor.set(qn('wp:layoutInCell'), '1')
    anchor.set(qn('wp:allowOverlap'), '1')

    simplePos = OxmlElement('wp:simplePos')
    simplePos.set('x', '0')
    simplePos.set('y', '0')
    anchor.append(simplePos)

    posH = OxmlElement('wp:positionH')
    posH.set(qn('wp:relativeFrom'), 'column')
    offH = OxmlElement('wp:posOffset')
    offH.text = '0'
    posH.append(offH)
    anchor.append(posH)

    posV = OxmlElement('wp:positionV')
    posV.set(qn('wp:relativeFrom'), 'paragraph')
    offV = OxmlElement('wp:posOffset')
    offV.text = '0'
    posV.append(offV)
    anchor.append(posV)

    # distances
    anchor.set(qn('wp:distT'), '0')
    anchor.set(qn('wp:distB'), '0')
    anchor.set(qn('wp:distL'), '0')
    anchor.set(qn('wp:distR'), '0')

    wrap = OxmlElement('wp:wrapNone')
    anchor.append(wrap)

    # move children from inline to anchor
    for child in list(wp_inline):
        anchor.append(child)

    drawing.remove(wp_inline)
    drawing.append(anchor)


# =========================
# DRAWINGML TEXTBOX (WPS) — SAFE FOR DOCX
# =========================
def _add_dml_textbox(paragraph, text: str, left_pt: float, top_pt: float,
                     width_pt: float, height_pt: float, shape_id: int):
    """
    Adds a WordprocessingShape textbox (DrawingML, wps:wsp) as a floating anchor.
    This is compatible with modern Word (.docx) and avoids VML corruption.
    """
    # w:r -> w:drawing -> wp:anchor -> a:graphic -> a:graphicData(uri=wps) -> wps:wsp ...
    r = paragraph.add_run()._r
    drawing = OxmlElement('w:drawing')

    anchor = OxmlElement('wp:anchor')
    anchor.set(qn('wp:simplePos'), '0')
    anchor.set(qn('wp:relativeHeight'), '251659265')  # above picture
    anchor.set(qn('wp:behindDoc'), '0')
    anchor.set(qn('wp:locked'), '0')
    anchor.set(qn('wp:layoutInCell'), '1')
    anchor.set(qn('wp:allowOverlap'), '1')
    anchor.set(qn('wp:distT'), '0')
    anchor.set(qn('wp:distB'), '0')
    anchor.set(qn('wp:distL'), '0')
    anchor.set(qn('wp:distR'), '0')

    simplePos = OxmlElement('wp:simplePos')
    simplePos.set('x', '0')
    simplePos.set('y', '0')
    anchor.append(simplePos)

    posH = OxmlElement('wp:positionH')
    posH.set(qn('wp:relativeFrom'), 'column')
    offH = OxmlElement('wp:posOffset')
    offH.text = str(_pt_to_emu(left_pt))
    posH.append(offH)
    anchor.append(posH)

    posV = OxmlElement('wp:positionV')
    posV.set(qn('wp:relativeFrom'), 'paragraph')
    offV = OxmlElement('wp:posOffset')
    offV.text = str(_pt_to_emu(top_pt))
    posV.append(offV)
    anchor.append(posV)

    extent = OxmlElement('wp:extent')
    extent.set('cx', str(_pt_to_emu(width_pt)))
    extent.set('cy', str(_pt_to_emu(height_pt)))
    anchor.append(extent)

    effect = OxmlElement('wp:effectExtent')
    effect.set('l', '0'); effect.set('t', '0'); effect.set('r', '0'); effect.set('b', '0')
    anchor.append(effect)

    wrap = OxmlElement('wp:wrapNone')
    anchor.append(wrap)

    docPr = OxmlElement('wp:docPr')
    docPr.set('id', str(shape_id))
    docPr.set('name', f'Label {shape_id}')
    anchor.append(docPr)

    cNv = OxmlElement('wp:cNvGraphicFramePr')
    gfLocks = OxmlElement('a:graphicFrameLocks')
    gfLocks.set(qn('xmlns:a'), 'http://schemas.openxmlformats.org/drawingml/2006/main')
    gfLocks.set('noChangeAspect', '1')
    cNv.append(gfLocks)
    anchor.append(cNv)

    graphic = OxmlElement('a:graphic')
    graphic.set(qn('xmlns:a'), 'http://schemas.openxmlformats.org/drawingml/2006/main')

    graphicData = OxmlElement('a:graphicData')
    graphicData.set('uri', 'http://schemas.microsoft.com/office/word/2010/wordprocessingShape')

    # wps:wsp root
    wsp = OxmlElement('wps:wsp')
    wsp.set(qn('xmlns:wps'), 'http://schemas.microsoft.com/office/word/2010/wordprocessingShape')
    wsp.set(qn('xmlns:wpg'), 'http://schemas.microsoft.com/office/word/2010/wordprocessingGroup')
    wsp.set(qn('xmlns:wpi'), 'http://schemas.microsoft.com/office/word/2010/wordprocessingInk')
    wsp.set(qn('xmlns:wpc'), 'http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas')
    wsp.set(qn('xmlns:a'), 'http://schemas.openxmlformats.org/drawingml/2006/main')
    wsp.set(qn('xmlns:r'), 'http://schemas.openxmlformats.org/officeDocument/2006/relationships')
    wsp.set(qn('xmlns:wp'), 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing')
    wsp.set(qn('xmlns:w'), 'http://schemas.openxmlformats.org/wordprocessingml/2006/main')

    # Non-visual properties
    cNvSpPr = OxmlElement('wps:cNvSpPr')
    cNvPr = OxmlElement('wps:cNvPr')
    cNvPr.set('id', str(shape_id))
    cNvPr.set('name', f'Label {shape_id}')
    cNvSpPr.append(cNvPr)
    cNvSpPr.append(OxmlElement('wps:cNvSpPr'))
    wsp.append(cNvSpPr)

    # Shape properties
    spPr = OxmlElement('wps:spPr')

    xfrm = OxmlElement('a:xfrm')
    off = OxmlElement('a:off'); off.set('x', '0'); off.set('y', '0')
    ext = OxmlElement('a:ext'); ext.set('cx', str(_pt_to_emu(width_pt))); ext.set('cy', str(_pt_to_emu(height_pt)))
    xfrm.append(off); xfrm.append(ext)
    spPr.append(xfrm)

    prst = OxmlElement('a:prstGeom'); prst.set('prst', 'rect')
    prst.append(OxmlElement('a:avLst'))
    spPr.append(prst)

    # Fill red
    fill = OxmlElement('a:solidFill')
    srgb = OxmlElement('a:srgbClr'); srgb.set('val', 'FF0000')
    fill.append(srgb)
    spPr.append(fill)

    # Line black 1.5pt
    ln = OxmlElement('a:ln'); ln.set('w', str(LINE_W_EMU))
    lnFill = OxmlElement('a:solidFill')
    lnRgb = OxmlElement('a:srgbClr'); lnRgb.set('val', '000000')
    lnFill.append(lnRgb)
    ln.append(lnFill)
    spPr.append(ln)

    wsp.append(spPr)

    # Textbox
    txbx = OxmlElement('wps:txbx')
    txbxContent = OxmlElement('w:txbxContent')

    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    jc = OxmlElement('w:jc'); jc.set(qn('w:val'), 'center')
    pPr.append(jc)
    p.append(pPr)

    rr = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    color = OxmlElement('w:color'); color.set(qn('w:val'), 'FFFFFF')
    b = OxmlElement('w:b')
    sz = OxmlElement('w:sz'); sz.set(qn('w:val'), '20')  # 10pt
    rPr.append(color); rPr.append(b); rPr.append(sz)
    rr.append(rPr)
    t = OxmlElement('w:t'); t.text = text
    rr.append(t)
    p.append(rr)

    txbxContent.append(p)
    txbx.append(txbxContent)
    wsp.append(txbx)

    bodyPr = OxmlElement('wps:bodyPr')
    bodyPr.set('wrap', 'none')
    bodyPr.set('lIns', str(_pt_to_emu(4)))
    bodyPr.set('rIns', str(_pt_to_emu(4)))
    bodyPr.set('tIns', str(_pt_to_emu(1)))
    bodyPr.set('bIns', str(_pt_to_emu(1)))
    # Auto-fit to text when possible
    bodyPr.append(OxmlElement('a:spAutoFit'))
    wsp.append(bodyPr)

    graphicData.append(wsp)
    graphic.append(graphicData)
    anchor.append(graphic)

    drawing.append(anchor)
    r.append(drawing)


def _add_two_labels(paragraph, width_cm: float, height_cm: float, shape_id_start: int):
    # Place at bottom-left and bottom-right of the image box (same paragraph anchor)
    w_pt = _cm_to_pt(width_cm)
    h_pt = _cm_to_pt(height_cm)
    top = h_pt - LABEL_H_PT - LABEL_PAD_PT

    left1 = LABEL_PAD_PT
    left2 = w_pt - LABEL_W_PT - LABEL_PAD_PT

    _add_dml_textbox(paragraph, LABEL_TEXT, left1, top, LABEL_W_PT, LABEL_H_PT, shape_id_start)
    _add_dml_textbox(paragraph, LABEL_TEXT, left2, top, LABEL_W_PT, LABEL_H_PT, shape_id_start + 1)


# =========================
# MAIN API (expected by your auto.py / auto_patched.py)
# =========================
def inserir_conteudo(modelo_path, conteudo, output_path):
    """
    Assinatura compatível com sua automação:
      inserir_conteudo(modelo_path, conteudo, output_path) -> contador_imagens
    """
    doc = Document(modelo_path)
    contador_imagens = 0
    paragrafo_insercao_index = None

    for i, paragrafo in enumerate(doc.paragraphs):
        if "{{start_here}}" in paragrafo.text:
            paragrafo_insercao_index = i
            break

    if paragrafo_insercao_index is None:
        print("⚠ Marca '{{start_here}}' não encontrada no modelo.")
        return contador_imagens

    conteudo_otimizado = otimizar_layout(conteudo)
    conteudo_invertido = list(reversed(conteudo_otimizado))

    shape_id = 9000  # ids únicos para docPr/cNvPr

    for item in conteudo_invertido:
        # Títulos
        if isinstance(item, str):
            titulo = item.replace("»", "").strip() + ":"
            nivel = item.count("»")
            p = doc.paragraphs[paragrafo_insercao_index].insert_paragraph_before("")
            run = p.add_run(titulo)

            if any(pasta in titulo for pasta in PASTAS_TEXTO_NORMAL):
                aplicar_estilo(run, 11, True)
                p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            elif nivel == 0:
                p.style = "Heading 1"
            elif nivel == 1:
                p.style = "Heading 2"
            elif nivel == 2:
                p.style = "Heading 3"
            else:
                aplicar_estilo(run, 12, True)

        # Imagem individual
        elif isinstance(item, dict) and "imagem" in item:
            imagem_path = item["imagem"]

            if os.path.exists(imagem_path) and os.path.getsize(imagem_path) > 0:
                try:
                    with Image.open(imagem_path) as img:
                        largura_original, altura_original = img.size

                        altura_desejada_cm = ALTURA_PADRAO
                        proporcao = altura_desejada_cm / (altura_original / 37.7952755906)
                        largura_proporcional_cm = largura_original * proporcao / 37.7952755906

                        p = doc.paragraphs[paragrafo_insercao_index].insert_paragraph_before("")
                        run = p.add_run()
                        run.add_picture(
                            imagem_path,
                            width=Cm(largura_proporcional_cm),
                            height=Cm(altura_desejada_cm)
                        )
                        _convert_last_inline_to_anchor(run)

                        if not _should_skip_labels(imagem_path):
                            _add_two_labels(p, largura_proporcional_cm, altura_desejada_cm, shape_id)
                            shape_id += 2

                        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        contador_imagens += 1

                except UnidentifiedImageError:
                    msg = f"[ERRO: formato não reconhecido] {imagem_path}"
                    print(msg)
                    p = doc.paragraphs[paragrafo_insercao_index].insert_paragraph_before(msg)
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                except Exception as e:
                    msg = f"[ERRO ao inserir imagem: {imagem_path}] {e}"
                    print(msg)
                    p = doc.paragraphs[paragrafo_insercao_index].insert_paragraph_before(msg)
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            else:
                msg = f"[ERRO: imagem inválida] {imagem_path}"
                print(msg)
                p = doc.paragraphs[paragrafo_insercao_index].insert_paragraph_before(msg)
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Tabela de imagens agrupadas
        elif isinstance(item, dict) and "tabela_imagens" in item:
            imagens = item["tabela_imagens"]
            colunas = item["colunas"]

            p_sep = doc.paragraphs[paragrafo_insercao_index].insert_paragraph_before("")
            tabela = doc.add_table(rows=1, cols=colunas)
            p_sep._p.addnext(tabela._tbl)
            tabela.autofit = False
            tabela.allow_autofit = False

            # remove borders
            for border in tabela._tbl.tblPr.xpath("./w:tblBorders"):
                border.getparent().remove(border)

            largura_coluna = 16.0 / colunas
            for col in tabela.columns:
                col.width = Cm(largura_coluna)

            for idx, img_path in enumerate(imagens):
                if idx < len(tabela.columns):
                    cell = tabela.cell(0, idx)
                    paragraph = cell.paragraphs[0]
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    run = paragraph.add_run()

                    try:
                        run.add_picture(img_path, height=Cm(ALTURA_PADRAO))
                        _convert_last_inline_to_anchor(run)

                        if not _should_skip_labels(img_path):
                            _add_two_labels(paragraph, largura_coluna, ALTURA_PADRAO, shape_id)
                            shape_id += 2

                        contador_imagens += 1
                    except Exception as e:
                        paragraph.add_run(f"[Erro img: {os.path.basename(img_path)}]")
                        print(f"Erro ao inserir na tabela: {e}")

        # Quebra de página
        elif isinstance(item, dict) and "quebra_pagina" in item:
            doc.paragraphs[paragrafo_insercao_index].insert_paragraph_before("").add_run().add_break(WD_BREAK.PAGE)

        # Parágrafo vazio
        elif isinstance(item, dict) and "paragrafo" in item:
            doc.paragraphs[paragrafo_insercao_index].insert_paragraph_before("")

    doc.save(output_path)
    print(f"✔ Imagens redimensionadas para 10 cm de altura ({contador_imagens} inseridas).")
    return contador_imagens
