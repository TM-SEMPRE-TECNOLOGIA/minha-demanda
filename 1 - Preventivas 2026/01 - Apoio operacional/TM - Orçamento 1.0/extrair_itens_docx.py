# Requisitos:
#   pip install python-docx openpyxl
#
# Uso:
#   python extrair_itens_docx.py
#
# O que faz:
# - Tkinter: seleciona .docx e escolhe onde salvar .xlsx
# - Lê SOMENTE tabelas cujo cabeçalho (1ª linha) contenha "Itens"
# - Extrai (Codigo, Quantidade), sem unidade
# - Gera um log (.txt) no mesmo diretório do .xlsx

import os
import re
from datetime import datetime
import tkinter as tk
from tkinter import filedialog, messagebox

from docx import Document
from openpyxl import Workbook
from openpyxl.utils import get_column_letter


CODE_RE = re.compile(r"^\s*\d+(?:\.\d+)?\s*$")  # aceita 17.4, 13.12, 2.24, 19.37 etc.


def norm(s: str) -> str:
    return (s or "").replace("\xa0", " ").strip()


def is_itens_table(table) -> bool:
    # Tabela-alvo: primeira linha contém "Itens" em qualquer célula
    if not table.rows:
        return False
    first_row_texts = [norm(c.text) for c in table.rows[0].cells]
    return any(t.lower() == "itens" for t in first_row_texts)


def pick_quantity_from_row(cells_text):
    """
    Preferência:
    - coluna 3 (índice 2), pois geralmente é a quantidade
    Fallback:
    - procurar o primeiro número no texto da linha (1,00 / 277,50 / 10 / 2,5 etc.)
    """
    if len(cells_text) >= 3:
        q = norm(cells_text[2])
        if q and q.upper() != "#N/D":
            return q

    joined = " ".join(cells_text)
    m = re.search(r"(\d{1,3}(?:\.\d{3})*,\d+|\d+,\d+|\d+)", joined)
    return m.group(1) if m else ""


def extract_code_qty(docx_path: str):
    doc = Document(docx_path)

    results = []
    logs = []
    itens_tables = 0

    for t_i, table in enumerate(doc.tables, start=1):
        if not is_itens_table(table):
            continue

        itens_tables += 1

        # pula cabeçalho "Itens"
        for r_i, row in enumerate(table.rows[1:], start=2):
            cells = [norm(c.text) for c in row.cells]
            if not cells:
                logs.append((t_i, r_i, "skip_empty_row", ""))
                continue

            code = norm(cells[0])
            if not code or code.upper() == "#N/D":
                logs.append((t_i, r_i, "skip_code_empty_or_ND", ""))
                continue

            if not CODE_RE.match(code):
                logs.append((t_i, r_i, "skip_code_invalid", code))
                continue

            qty = pick_quantity_from_row(cells)
            if not qty or qty.upper() == "#N/D":
                logs.append((t_i, r_i, "skip_qty_empty_or_ND", code))
                continue

            results.append((code.strip(), qty.strip()))

    # Não remover duplicados - o usuário quer todas as ocorrências
    # (Removida lógica de set/seen)

    meta = {
        "tables_total": len(doc.tables),
        "itens_tables": itens_tables,
        "rows_extracted": len(results),
        "rows_ignored": len(logs),
        "ignored_details": logs
    }
    return results, meta


def save_xlsx(rows, out_path: str):
    wb = Workbook()
    ws = wb.active
    ws.title = "Itens"

    ws.append(["Codigo", "Quantidade"])
    for code, qty in rows:
        ws.append([code, qty])

    ws.freeze_panes = "A2"

    # largura automática simples
    for col in range(1, 3):
        max_len = 0
        for cell in ws[get_column_letter(col)]:
            v = "" if cell.value is None else str(cell.value)
            max_len = max(max_len, len(v))
        ws.column_dimensions[get_column_letter(col)].width = min(max_len + 2, 40)

    wb.save(out_path)


def save_log(meta, rows, docx_path, out_xlsx_path):
    log_path = os.path.splitext(out_xlsx_path)[0] + "_log.txt"

    with open(log_path, "w", encoding="utf-8") as f:
        f.write("LOG - Extração de Tabelas 'Itens'\n")
        f.write(f"Arquivo: {os.path.basename(docx_path)}\n")
        f.write(f"Data/hora: {datetime.now().isoformat(timespec='seconds')}\n\n")

        f.write(f"Tabelas totais no DOCX: {meta['tables_total']}\n")
        f.write(f"Tabelas identificadas como 'Itens': {meta['itens_tables']}\n")
        f.write(f"Linhas extraídas (total): {meta['rows_extracted']}\n")
        f.write(f"Linhas ignoradas: {meta['rows_ignored']}\n\n")

        if meta["ignored_details"]:
            f.write("Detalhes ignorados (tabela, linha, motivo, valor):\n")
            for t_i, r_i, reason, val in meta["ignored_details"]:
                f.write(f"- T{t_i} L{r_i}: {reason} {val}\n")
            f.write("\n")

        f.write("Saída (Codigo | Quantidade):\n")
        for c, q in rows:
            f.write(f"{c} | {q}\n")

    return log_path


def run_gui():
    root = tk.Tk()
    root.withdraw()

    docx_path = filedialog.askopenfilename(
        title="Selecione o arquivo Word (.docx)",
        filetypes=[("Word (.docx)", "*.docx")]
    )
    if not docx_path:
        return

    out_xlsx = filedialog.asksaveasfilename(
        title="Salvar saída em planilha (.xlsx)",
        defaultextension=".xlsx",
        filetypes=[("Excel (.xlsx)", "*.xlsx")]
    )
    if not out_xlsx:
        return

    try:
        rows, meta = extract_code_qty(docx_path)

        if not rows:
            messagebox.showwarning(
                "Sem dados",
                "Nenhuma linha válida foi encontrada nas tabelas 'Itens'."
            )
            return

        save_xlsx(rows, out_xlsx)
        log_path = save_log(meta, rows, docx_path, out_xlsx)

        messagebox.showinfo(
            "Concluído",
            f"Planilha criada:\n{out_xlsx}\n\nLog criado:\n{log_path}"
        )
    except Exception as e:
        messagebox.showerror("Erro", f"Ocorreu um erro:\n{e}")


if __name__ == "__main__":
    run_gui()
